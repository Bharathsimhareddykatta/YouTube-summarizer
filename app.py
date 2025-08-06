import streamlit as st
import os
from utils.youtube_utils import get_transcript, test_video_accessibility, clean_transcript
from utils.summarizer import summarize_with_api
from fpdf import FPDF
from docx import Document

# Initialize session state
if 'transcript' not in st.session_state:
    st.session_state.transcript = ""
if 'summary' not in st.session_state:
    st.session_state.summary = ""

# --- Export Utilities ---
def export_as_pdf(summary):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in summary.split('\n'):
        pdf.multi_cell(0, 10, line)
    pdf_path = "summary.pdf"
    pdf.output(pdf_path)
    return pdf_path

def export_as_docx(summary):
    doc = Document()
    doc.add_heading("YouTube Summary", 0)
    doc.add_paragraph(summary)
    doc_path = "summary.docx"
    doc.save(doc_path)
    return doc_path

# --- Streamlit UI ---
st.set_page_config(
    page_title="YouTube Script Summarizer", 
    page_icon="🎬",
    layout="wide"
)

st.title("🎬 YouTube Script Summarizer")
st.markdown("Extract and summarize YouTube video transcripts using AI")

# API Key Check
api_key = os.getenv("OPENROUTER_API_KEY")
if not api_key:
    st.error("❌ OPENROUTER_API_KEY not found. Please set your API key in the environment variables.")
    st.info("💡 **To set your API key:**")
    st.markdown("""
    1. Get your API key from [OpenRouter](https://openrouter.ai/keys)
    2. Set it as an environment variable: `OPENROUTER_API_KEY=your_key_here`
    3. Restart the application
    """)
    st.stop()

# Main content
input_mode = st.radio("Choose Input Type", ["YouTube URL", "Paste Transcript"])

if input_mode == "YouTube URL":
    video_url = st.text_input("Enter YouTube Video URL:")
    if st.button("🔍 Fetch Transcript", type="primary"):
        if video_url.strip():
            with st.spinner("Fetching and cleaning transcript..."):
                transcript_result = get_transcript(video_url)
                if transcript_result.startswith("❌"):
                    st.error(transcript_result)
                    st.info("💡 **Troubleshooting Tips:**")
                    st.markdown("""
                    - Make sure the video has captions/transcripts enabled
                    - Try a different YouTube video
                    - Check your internet connection
                    - Some videos may have restricted transcript access
                    """)
                else:
                    st.session_state.transcript = transcript_result
                    st.success("✅ Transcript fetched and cleaned successfully!")
                    st.rerun()
        else:
            st.warning("Please enter a YouTube URL")
else:
    manual_transcript = st.text_area("Paste Transcript", height=300, placeholder="Paste your transcript here...")
    if st.button("📝 Save Transcript"):
        if manual_transcript.strip():
            st.session_state.transcript = clean_transcript(manual_transcript)
            st.success("✅ Transcript saved and cleaned!")
            st.rerun()
        else:
            st.warning("Please paste a transcript")

# Display current transcript
if st.session_state.transcript:
    st.divider()
    st.subheader("📝 Transcript")
    
    st.text_area("Cleaned Transcript (Ready for Summarization)", 
                st.session_state.transcript, height=300, key="display_transcript")
    st.info(f"📊 Length: {len(st.session_state.transcript)} characters")
    
    # Summarization button
    if st.button("📊 Generate AI Summary", type="primary"):
        with st.spinner("Generating AI summary..."):
            try:
                summary_result = summarize_with_api(st.session_state.transcript)
                
                if summary_result.startswith("❌"):
                    st.error(summary_result)
                    st.info("💡 **API Troubleshooting:**")
                    st.markdown("""
                    - **Check API Key**: Make sure OPENROUTER_API_KEY is set correctly
                    - **Verify Credits**: Check your OpenRouter account has sufficient credits
                    - **Check Internet**: Ensure you have a stable internet connection
                    """)
                else:
                    st.session_state.summary = summary_result
                    st.rerun()
            except Exception as e:
                st.error(f"Summarization failed: {str(e)}")

# Display summary
if st.session_state.summary:
    st.divider()
    st.subheader("🤖 AI Summary")
    st.write(st.session_state.summary)
    
    # Export options
    st.divider()
    st.subheader("💾 Export Options")
    
    col1, col2 = st.columns(2)
    with col1:
        try:
            pdf_path = export_as_pdf(st.session_state.summary)
            with open(pdf_path, "rb") as f:
                st.download_button(
                    "📄 Download PDF", 
                    data=f.read(), 
                    file_name="youtube_summary.pdf",
                    mime="application/pdf"
                )
        except Exception as e:
            st.error(f"PDF export failed: {str(e)}")
    
    with col2:
        try:
            docx_path = export_as_docx(st.session_state.summary)
            with open(docx_path, "rb") as f:
                st.download_button(
                    "📝 Download DOCX", 
                    data=f.read(), 
                    file_name="youtube_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"DOCX export failed: {str(e)}")

# Footer
st.divider()
st.markdown("---")
st.markdown(
    """
    <div style='text-align: center; color: #666;'>
        <p>Made with ❤️ using Streamlit & OpenRouter AI</p>
        <p>Powered by Claude AI for intelligent summarization</p>
    </div>
    """,
    unsafe_allow_html=True
)
